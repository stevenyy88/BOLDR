"""
BOLDR Self-Improving Customer Intelligence Engine
Knowledge Base Indexer — Parse all KB source files into structured documents

Parses 5 source documents:
  1. 02_product_reference.docx — model specs, strap catalogue, quick answers
  2. 03a_rate_card_engraving.csv — engraving pricing and turnaround
  3. 03b_rate_card_servicing.csv — servicing tiers and pricing
  4. 04_faq_document.pdf — FAQ entries across 6 categories
  5. 05_cs_sop.docx — escalation rules, tone guidelines, enquiry handling, new questions log

Produces a list of KBDocument-like dicts ready for ChromaDB insertion.

Author: Steve Ng, Founder and CEO - Digital Futures Consultancy LLP
"""

import csv
import logging
import re
from dataclasses import asdict
from pathlib import Path
from typing import Any

import pandas as pd
import PyPDF2
from docx import Document

from app.kb.schemas import (
    EngravingService,
    FAQEntry,
    KBDocument,
    KBSource,
    ModelSpec,
    NewQuestionLogEntry,
    ServicingTier,
    SOPEscalationRule,
    StrapEntry,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Category constants (used in metadata for filtered retrieval)
# ---------------------------------------------------------------------------
CAT_PRODUCT_SPECS = "product_specs"
CAT_STRAP_CATALOGUE = "strap_catalogue"
CAT_QUICK_ANSWERS = "quick_answers"
CAT_ENGRAVING = "engraving"
CAT_SERVICING = "servicing"
CAT_FAQ = "faq"
CAT_SOP = "sop"
CAT_SOP_ESCALATION = "sop_escalation"
CAT_SOP_TONE = "sop_tone"
CAT_SOP_ENQUIRY = "sop_enquiry"
CAT_SOP_NEW_QUESTIONS = "sop_new_questions"
CAT_SOP_CONTACTS = "sop_contacts"

# FAQ categories extracted from the PDF structure
FAQ_CATEGORIES = [
    "Materials & Safety",
    "Engraving",
    "Strap Compatibility",
    "Watch Servicing",
    "Orders & Shipping",
    "General",
]


class KBIndexer:
    """Parse all BOLDR knowledge-base source files into structured documents.

    Each ``parse_*`` method returns a list of dicts (or a list of dataclass
    instances) that can be freely converted to ``KBDocument`` objects and
    subsequently indexed into ChromaDB.

    Usage::

        indexer = KBIndexer(data_dir="/path/to/dataset")
        documents = indexer.index_all()
    """

    def __init__(self, data_dir: str | Path | None = None):
        if data_dir is None:
            data_dir = Path(__file__).resolve().parent.parent.parent / "dataset"
        self.data_dir = Path(data_dir)

        # File paths
        self.product_ref_path = self.data_dir / "02_product_reference.docx"
        self.engraving_csv_path = self.data_dir / "03a_rate_card_engraving.csv"
        self.servicing_csv_path = self.data_dir / "03b_rate_card_servicing.csv"
        self.faq_pdf_path = self.data_dir / "04_faq_document.pdf"
        self.cs_sop_path = self.data_dir / "05_cs_sop.docx"

        logger.info("KBIndexer initialised — data_dir=%s", self.data_dir)

    # ------------------------------------------------------------------
    # 1. Product Reference (DOCX)
    # ------------------------------------------------------------------

    def parse_product_reference(self) -> list[dict]:
        """Parse 02_product_reference.docx.

        Extracts:
        - 3 model spec tables  → individual spec documents per model
        - Strap catalogue table → one document per strap SKU
        - Quick-answer table    → one document per Q&A pair
        """
        path = self.product_ref_path
        if not path.exists():
            logger.error("Product reference file not found: %s", path)
            return []

        logger.info("Parsing product reference: %s", path)
        doc = Document(str(path))
        documents: list[dict] = []

        # --- Model headings (for context) ---
        model_headings: dict[int, str] = {}
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if "Heading 2" in style_name and text:
                # Extract model name from heading like "Expedition Titanium  |  SGD 485  |  SKU: BLD-EXP-TI-40"
                model_headings[len(model_headings)] = text

        # --- Tables ---
        tables = doc.tables
        if len(tables) < 5:
            logger.warning(
                "Expected at least 5 tables in product reference, found %d", len(tables)
            )

        # Tables 0–2: Model spec tables
        model_names = [
            "Expedition Titanium",
            "Journey Titanium",
            "Expedition Titanium — Ember Limited Edition",
        ]
        model_skus = ["BLD-EXP-TI-40", "BLD-JRN-TI-38", "BLD-EXP-TI-40-LE"]
        model_prices = ["SGD 485", "SGD 395", "SGD 595"]

        for idx in range(min(3, len(tables))):
            table = tables[idx]
            name = model_names[idx] if idx < len(model_names) else f"Model_{idx}"
            sku = model_skus[idx] if idx < len(model_skus) else ""
            price = model_prices[idx] if idx < len(model_prices) else ""

            specs: dict[str, str] = {}
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    specs[cells[0]] = cells[1]

            # Build rich content string for embedding
            content_parts = [f"BOLDR {name} (SKU: {sku}, Price: {price})"]
            for key, val in specs.items():
                content_parts.append(f"{key}: {val}")
            content = "\n".join(content_parts)

            doc_id = f"product_spec_{sku.lower().replace('-', '_')}"

            documents.append(
                {
                    "id": doc_id,
                    "content": content,
                    "metadata": {
                        "model_name": name,
                        "sku": sku,
                        "price": price,
                    },
                    "source_file": KBSource.PRODUCT_REFERENCE.value,
                    "section": "model_specs",
                    "category": CAT_PRODUCT_SPECS,
                }
            )
            logger.debug("Indexed model spec: %s", name)

        # Table 3: Strap catalogue
        if len(tables) >= 4:
            strap_table = tables[3]
            headers = [c.text.strip() for c in strap_table.rows[0].cells]
            for row_idx, row in enumerate(strap_table.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if not cells[0]:
                    continue

                row_dict = dict(zip(headers, cells))

                content = (
                    f"BOLDR Strap — {row_dict.get('Type', '')} "
                    f"({row_dict.get('Colour', '')})\n"
                    f"SKU: {row_dict.get('SKU', '')}\n"
                    f"Price: {row_dict.get('Price (SGD)', '')}\n"
                    f"BPA-free: {row_dict.get('BPA-free', '')}\n"
                    f"Compatible with: {row_dict.get('Compatible with', '')}"
                )

                doc_id = f"strap_{row_dict.get('SKU', row_idx).lower().replace('-', '_')}"

                documents.append(
                    {
                        "id": doc_id,
                        "content": content,
                        "metadata": {
                            "sku": row_dict.get("SKU", ""),
                            "type": row_dict.get("Type", ""),
                            "colour": row_dict.get("Colour", ""),
                            "bpa_free": row_dict.get("BPA-free", ""),
                            "price": row_dict.get("Price (SGD)", ""),
                            "compatible_with": row_dict.get("Compatible with", ""),
                        },
                        "source_file": KBSource.PRODUCT_REFERENCE.value,
                        "section": "strap_catalogue",
                        "category": CAT_STRAP_CATALOGUE,
                    }
                )
            logger.info("Indexed %d strap entries", len([d for d in documents if d["category"] == CAT_STRAP_CATALOGUE]))

        # Table 4: Quick answers
        if len(tables) >= 5:
            qa_table = tables[4]
            for row_idx, row in enumerate(qa_table.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    question = cells[0].strip('"').strip()
                    answer = cells[1]

                    content = f"Q: {question}\nA: {answer}"

                    doc_id = f"quick_answer_{row_idx:02d}"

                    documents.append(
                        {
                            "id": doc_id,
                            "content": content,
                            "metadata": {
                                "question": question,
                                "answer": answer,
                            },
                            "source_file": KBSource.PRODUCT_REFERENCE.value,
                            "section": "quick_answers",
                            "category": CAT_QUICK_ANSWERS,
                        }
                    )
            logger.info("Indexed %d quick-answer entries", len([d for d in documents if d["category"] == CAT_QUICK_ANSWERS]))

        logger.info("Product reference: %d total documents", len(documents))
        return documents

    # ------------------------------------------------------------------
    # 2. Engraving Rate Card (CSV)
    # ------------------------------------------------------------------

    def parse_engraving_rate_card(self) -> list[dict]:
        """Parse 03a_rate_card_engraving.csv.

        Each row becomes one document with service, price, and notes.
        """
        path = self.engraving_csv_path
        if not path.exists():
            logger.error("Engraving rate card not found: %s", path)
            return []

        logger.info("Parsing engraving rate card: %s", path)
        df = pd.read_csv(path)
        documents: list[dict] = []

        for idx, row in df.iterrows():
            service = str(row.get("service", "")).strip()
            price = row.get("price_sgd", 0)
            notes = str(row.get("notes", "")).strip()

            content = f"BOLDR Engraving Service: {service}\nPrice: SGD {price}\nNotes: {notes}"

            # Create a URL-safe doc ID
            safe_id = re.sub(r"[^a-z0-9]+", "_", service.lower())[:60]
            doc_id = f"engraving_{safe_id}"

            documents.append(
                {
                    "id": doc_id,
                    "content": content,
                    "metadata": {
                        "service": service,
                        "price_sgd": float(price),
                        "notes": notes,
                    },
                    "source_file": KBSource.RATE_CARD_ENGRAVING.value,
                    "section": "engraving_pricing",
                    "category": CAT_ENGRAVING,
                }
            )

        logger.info("Engraving rate card: %d documents", len(documents))
        return documents

    # ------------------------------------------------------------------
    # 3. Servicing Rate Card (CSV)
    # ------------------------------------------------------------------

    def parse_servicing_rate_card(self) -> list[dict]:
        """Parse 03b_rate_card_servicing.csv.

        Each row becomes one document with service tier, price, turnaround,
        inclusions, and notes.
        """
        path = self.servicing_csv_path
        if not path.exists():
            logger.error("Servicing rate card not found: %s", path)
            return []

        logger.info("Parsing servicing rate card: %s", path)
        df = pd.read_csv(path)
        documents: list[dict] = []

        for idx, row in df.iterrows():
            tier = str(row.get("service_tier", "")).strip()
            price = row.get("price_sgd", 0)
            turnaround = str(row.get("turnaround_days", "")).strip()
            includes = str(row.get("includes", "")).strip()
            notes = str(row.get("notes", "")).strip()

            content_parts = [
                f"BOLDR Servicing: {tier}",
                f"Price: SGD {price}",
                f"Turnaround: {turnaround}",
                f"Includes: {includes}",
            ]
            if notes and notes.lower() != "nan":
                content_parts.append(f"Notes: {notes}")
            content = "\n".join(content_parts)

            safe_id = re.sub(r"[^a-z0-9]+", "_", tier.lower())[:60]
            doc_id = f"servicing_{safe_id}"

            metadata: dict[str, Any] = {
                "service_tier": tier,
                "price_sgd": float(price),
                "turnaround": turnaround,
                "includes": includes,
            }
            if notes and notes.lower() != "nan":
                metadata["notes"] = notes

            documents.append(
                {
                    "id": doc_id,
                    "content": content,
                    "metadata": metadata,
                    "source_file": KBSource.RATE_CARD_SERVICING.value,
                    "section": "servicing_pricing",
                    "category": CAT_SERVICING,
                }
            )

        logger.info("Servicing rate card: %d documents", len(documents))
        return documents

    # ------------------------------------------------------------------
    # 4. FAQ Document (PDF)
    # ------------------------------------------------------------------

    def parse_faq_document(self) -> list[dict]:
        """Parse 04_faq_document.pdf.

        Extracts all FAQ entries, categorised by section heading.
        The PDF has 6 categories and 28 Q&A pairs total.
        """
        path = self.faq_pdf_path
        if not path.exists():
            logger.error("FAQ document not found: %s", path)
            return []

        logger.info("Parsing FAQ document: %s", path)
        reader = PyPDF2.PdfReader(str(path))

        # Extract full text from all pages
        full_text = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            full_text += page_text + "\n"

        # Parse Q&A pairs with category detection
        documents: list[dict] = []
        current_category = "General"
        faq_index = 0

        # Split by lines and process
        lines = full_text.split("\n")

        # Category heading detection: these are section headers in the PDF
        category_keywords = {
            "Materials & Safety": "Materials & Safety",
            "Engraving": "Engraving",
            "Strap Compatibility": "Strap Compatibility",
            "Watch Servicing": "Watch Servicing",
            "Orders & Shipping": "Orders & Shipping",
            "General": "General",
        }

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Check if this line is a category heading (standalone, not a Q&A)
            is_category_heading = False
            for keyword, cat_name in category_keywords.items():
                if stripped == keyword and not stripped.startswith("Q:"):
                    current_category = cat_name
                    is_category_heading = True
                    break

            if is_category_heading:
                continue

            # Check if this line contains a Q&A pair
            # Format: "Q: ... A: ..."
            q_match = re.match(r"^Q:\s*(.+?)\s+A:\s*(.+)$", stripped, re.DOTALL)
            if q_match:
                faq_index += 1
                question = q_match.group(1).strip()
                answer = q_match.group(2).strip()

                content = f"FAQ ({current_category})\nQ: {question}\nA: {answer}"

                # Create a safe category ID
                cat_id = re.sub(r"[^a-z0-9]+", "_", current_category.lower())

                doc_id = f"faq_{cat_id}_{faq_index:02d}"

                documents.append(
                    {
                        "id": doc_id,
                        "content": content,
                        "metadata": {
                            "category": current_category,
                            "question": question,
                            "answer": answer,
                        },
                        "source_file": KBSource.FAQ_DOCUMENT.value,
                        "section": f"faq_{cat_id}",
                        "category": CAT_FAQ,
                    }
                )

        logger.info("FAQ document: %d documents across categories", len(documents))
        return documents

    # ------------------------------------------------------------------
    # 5. CS SOP (DOCX)
    # ------------------------------------------------------------------

    def parse_cs_sop(self) -> list[dict]:
        """Parse 05_cs_sop.docx.

        Extracts:
        - Reference document table (Section 2)
        - Enquiry handling steps (Section 3)
        - Enquiry type handling rules (Section 4)
        - Tone & brand voice guidelines (Section 5)
        - New questions log table (Section 6)
        - Escalation rules (Section 7)
        - Contacts table (Section 8)
        """
        path = self.cs_sop_path
        if not path.exists():
            logger.error("CS SOP file not found: %s", path)
            return []

        logger.info("Parsing CS SOP: %s", path)
        doc = Document(str(path))
        documents: list[dict] = []

        # --- Extract structured sections from paragraphs ---
        # Build a section map from paragraphs
        sections: dict[str, list[str]] = {}
        current_heading = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            is_heading = "Heading" in style_name

            if is_heading:
                # Extract section number and title
                # e.g. "1. Overview", "4.2 Engraving", etc.
                current_heading = text
                if current_heading not in sections:
                    sections[current_heading] = []
            else:
                if current_heading:
                    sections[current_heading].append(text)

        # --- Section 3: Step-by-step enquiry handling ---
        section_3_key = None
        for key in sections:
            if key.startswith("3."):
                section_3_key = key
                break

        if section_3_key and sections.get(section_3_key):
            steps = sections[section_3_key]
            content = "CS SOP — Enquiry Handling Steps\n\n"
            for i, step in enumerate(steps, 1):
                content += f"Step {i}: {step}\n"

            documents.append(
                {
                    "id": "sop_enquiry_handling",
                    "content": content,
                    "metadata": {"section": "3", "type": "enquiry_handling_steps"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "enquiry_handling",
                    "category": CAT_SOP_ENQUIRY,
                }
            )

        # --- Section 4: Enquiry type handling rules ---
        # Each subsection (4.1–4.6) becomes its own document
        enquiry_subsections = {}
        for key, lines in sections.items():
            if re.match(r"^4\.\d", key):
                enquiry_subsections[key] = lines

        for key, lines in enquiry_subsections.items():
            content_parts = [f"CS SOP — Enquiry Type: {key}"]
            content_parts.extend(lines)
            content = "\n".join(content_parts)

            # Extract subsection slug from heading
            slug = re.sub(r"[^a-z0-9]+", "_", key.lower())[:40]
            doc_id = f"sop_enquiry_{slug}"

            # Try to extract the subsection title for metadata
            subsection_title = key.split(" ", 1)[-1] if " " in key else key

            documents.append(
                {
                    "id": doc_id,
                    "content": content,
                    "metadata": {
                        "section": key,
                        "subsection_title": subsection_title,
                        "type": "enquiry_handling",
                    },
                    "source_file": KBSource.CS_SOP.value,
                    "section": f"enquiry_{slug}",
                    "category": CAT_SOP_ENQUIRY,
                }
            )

        # --- Section 5: Tone & Brand Voice ---
        section_5_key = None
        for key in sections:
            if key.startswith("5."):
                section_5_key = key
                break

        if section_5_key and sections.get(section_5_key):
            lines = sections[section_5_key]
            content = "CS SOP — Tone & Brand Voice\n\n"
            content += "\n".join(lines)

            documents.append(
                {
                    "id": "sop_tone_guidelines",
                    "content": content,
                    "metadata": {"section": "5", "type": "tone_guidelines"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "tone_guidelines",
                    "category": CAT_SOP_TONE,
                }
            )

        # --- Section 7: Escalation rules ---
        section_7_key = None
        for key in sections:
            if key.startswith("7."):
                section_7_key = key
                break

        if section_7_key and sections.get(section_7_key):
            lines = sections[section_7_key]
            content = "CS SOP — Escalation Rules\n\nWhen to escalate:\n"
            for line in lines:
                content += f"• {line}\n"

            documents.append(
                {
                    "id": "sop_escalation_rules",
                    "content": content,
                    "metadata": {"section": "7", "type": "escalation_rules"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "escalation",
                    "category": CAT_SOP_ESCALATION,
                }
            )

        # --- Table 0: Reference documents ---
        if len(doc.tables) >= 1:
            ref_table = doc.tables[0]
            content = "CS SOP — Reference Documents\n\n"
            for row in ref_table.rows[1:]:  # skip header
                cells = [c.text.strip() for c in row.cells]
                if cells[0]:
                    content += f"• {cells[0]}: {cells[1]} (Location: {cells[2]})\n"

            documents.append(
                {
                    "id": "sop_reference_documents",
                    "content": content,
                    "metadata": {"section": "2", "type": "reference_documents"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "reference_documents",
                    "category": CAT_SOP,
                }
            )

        # --- Table 1: Servicing quick-reference (in SOP) ---
        if len(doc.tables) >= 2:
            service_table = doc.tables[1]
            headers = [c.text.strip() for c in service_table.rows[0].cells]
            content = "CS SOP — Servicing Quick Reference\n\n"
            for row in service_table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if cells[0]:
                    row_dict = dict(zip(headers, cells))
                    parts = [f"{h}: {v}" for h, v in row_dict.items() if v]
                    content += "• " + " | ".join(parts) + "\n"

            documents.append(
                {
                    "id": "sop_servicing_quick_ref",
                    "content": content,
                    "metadata": {"section": "4.4", "type": "servicing_quick_reference"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "servicing_quick_reference",
                    "category": CAT_SOP,
                }
            )

        # --- Table 2: New Questions Log ---
        if len(doc.tables) >= 3:
            log_table = doc.tables[2]
            headers = [c.text.strip() for c in log_table.rows[0].cells]
            log_idx = 0
            for row in log_table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if not cells[0]:
                    continue  # skip empty rows

                row_dict = dict(zip(headers, cells))
                log_idx += 1

                content = (
                    f"New Question Log — {row_dict.get('Customer question (paraphrase)', '')}\n"
                    f"Date: {row_dict.get('Date', '')}\n"
                    f"Theme: {row_dict.get('Theme', '')}\n"
                    f"Status: {row_dict.get('Status', '')}"
                )

                slug = re.sub(r"[^a-z0-9]+", "_", row_dict.get("Theme", "").lower())[:30]
                doc_id = f"sop_new_question_{log_idx:02d}_{slug}"

                documents.append(
                    {
                        "id": doc_id,
                        "content": content,
                        "metadata": {
                            "date": row_dict.get("Date", ""),
                            "theme": row_dict.get("Theme", ""),
                            "status": row_dict.get("Status", ""),
                            "section": "6",
                        },
                        "source_file": KBSource.CS_SOP.value,
                        "section": "new_questions_log",
                        "category": CAT_SOP_NEW_QUESTIONS,
                    }
                )

        # --- Table 3: Contacts ---
        if len(doc.tables) >= 4:
            contact_table = doc.tables[3]
            content = "CS SOP — Contact Directory\n\n"
            for row in contact_table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if cells[0]:
                    content += f"• {cells[0]}: {cells[1]} — {cells[2]}\n"

            documents.append(
                {
                    "id": "sop_contacts",
                    "content": content,
                    "metadata": {"section": "8", "type": "contacts"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "contacts",
                    "category": CAT_SOP_CONTACTS,
                }
            )

        # --- Section 6: New Questions Log procedure ---
        section_6_key = None
        for key in sections:
            if key.startswith("6."):
                section_6_key = key
                break

        if section_6_key and sections.get(section_6_key):
            lines = sections[section_6_key]
            content = "CS SOP — New Questions Log Procedure\n\n"
            for i, line in enumerate(lines, 1):
                content += f"{i}. {line}\n"

            documents.append(
                {
                    "id": "sop_new_questions_procedure",
                    "content": content,
                    "metadata": {"section": "6", "type": "new_questions_procedure"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "new_questions_procedure",
                    "category": CAT_SOP_NEW_QUESTIONS,
                }
            )

        # --- Section 1: Overview ---
        section_1_key = None
        for key in sections:
            if key.startswith("1."):
                section_1_key = key
                break

        if section_1_key and sections.get(section_1_key):
            lines = sections[section_1_key]
            content = "CS SOP — Overview\n\n" + "\n".join(lines)

            documents.append(
                {
                    "id": "sop_overview",
                    "content": content,
                    "metadata": {"section": "1", "type": "overview"},
                    "source_file": KBSource.CS_SOP.value,
                    "section": "overview",
                    "category": CAT_SOP,
                }
            )

        logger.info("CS SOP: %d documents", len(documents))
        return documents

    # ------------------------------------------------------------------
    # Master indexer
    # ------------------------------------------------------------------

    def index_all(self) -> list[dict]:
        """Parse all source files and return a unified list of document dicts.

        Each dict has keys: id, content, metadata, source_file, section, category
        — matching the KBDocument schema.
        """
        all_documents: list[dict] = []

        # 1. Product reference
        try:
            docs = self.parse_product_reference()
            all_documents.extend(docs)
            logger.info("✓ Product reference: %d documents", len(docs))
        except Exception:
            logger.exception("Failed to parse product reference")

        # 2. Engraving rate card
        try:
            docs = self.parse_engraving_rate_card()
            all_documents.extend(docs)
            logger.info("✓ Engraving rate card: %d documents", len(docs))
        except Exception:
            logger.exception("Failed to parse engraving rate card")

        # 3. Servicing rate card
        try:
            docs = self.parse_servicing_rate_card()
            all_documents.extend(docs)
            logger.info("✓ Servicing rate card: %d documents", len(docs))
        except Exception:
            logger.exception("Failed to parse servicing rate card")

        # 4. FAQ document
        try:
            docs = self.parse_faq_document()
            all_documents.extend(docs)
            logger.info("✓ FAQ document: %d documents", len(docs))
        except Exception:
            logger.exception("Failed to parse FAQ document")

        # 5. CS SOP
        try:
            docs = self.parse_cs_sop()
            all_documents.extend(docs)
            logger.info("✓ CS SOP: %d documents", len(docs))
        except Exception:
            logger.exception("Failed to parse CS SOP")

        logger.info("=" * 50)
        logger.info("Total documents indexed: %d", len(all_documents))

        # Validate — every doc must have required keys
        required_keys = {"id", "content", "metadata", "source_file", "section", "category"}
        for doc in all_documents:
            missing = required_keys - set(doc.keys())
            if missing:
                logger.warning(
                    "Document %s missing keys: %s", doc.get("id", "UNKNOWN"), missing
                )

        return all_documents


# ---------------------------------------------------------------------------
# Main entry point (standalone)
# ---------------------------------------------------------------------------

def main() -> list[dict]:
    """Index all KB documents and return them.

    This is the entry point called by ``scripts/index_kb.py``.
    """
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )

    indexer = KBIndexer()
    documents = indexer.index_all()

    print(f"\n{'=' * 60}")
    print(f"BOLDR KB Indexer — {len(documents)} documents generated")
    print(f"{'=' * 60}")

    # Summary by source
    from collections import Counter

    sources = Counter(d["source_file"] for d in documents)
    print("\n📄 By source file:")
    for source, count in sources.most_common():
        print(f"   {source}: {count} documents")

    categories = Counter(d["category"] for d in documents)
    print("\n🏷️  By category:")
    for cat, count in categories.most_common():
        print(f"   {cat}: {count} documents")

    return documents


if __name__ == "__main__":
    docs = main()
    print(f"\n✅ Done — {len(docs)} total documents ready for ChromaDB")